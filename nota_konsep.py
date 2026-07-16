from __future__ import annotations

import os
import re
import tempfile
from datetime import datetime
from pathlib import Path
from typing import Iterable
from xml.sax.saxutils import escape as xml_escape

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas
from reportlab.platypus import Paragraph

MONTH_NAMES_ID = {
    1: "Januari", 2: "Februari", 3: "Maret", 4: "April",
    5: "Mei", 6: "Juni", 7: "Juli", 8: "Agustus",
    9: "September", 10: "Oktober", 11: "November", 12: "Desember",
}

HEADER_LINES = [
    "DINAS KEBUDAYAAN DAN PARIWISATA",
    "PROVINSI JAWA TIMUR",
    "BIDANG PEMASARAN",
]
FIXED_ADDRESSEE = [
    "Kepada",
    "Yth. Ibu Kepala Dinas Kebudayaan",
    "    dan Pariwisata Provinsi Jawa Timur",
    "    di",
    "        SURABAYA",
]
FROM_TEXT = "Bidang Pemasaran dan Kelembagaan Parekraf"
CLOSING_TEXT = (
    "Untuk mohon tanda tangan Ibu Kepala Dinas Kebudayaan dan Pariwisata "
    "Provinsi Jawa Timur."
)
SIGNATURE_TITLE = ["KEPALA BIDANG PEMASARAN DAN", "KELEMBAGAAN PAREKRAF"]
SIGNATURE_NAME = "ALI AFANDI, S.Pd, MT"
SIGNATURE_RANK = "Pembina Tk I"
SIGNATURE_NIP = "NIP. 197403252000031003"


def format_date_long_id(value: str) -> str:
    dt = datetime.strptime(value, "%Y-%m-%d")
    return f"{dt.day} {MONTH_NAMES_ID[dt.month]} {dt.year}"


def clean_recipients(raw: str | Iterable[str]) -> list[str]:
    if isinstance(raw, str):
        parts = re.split(r"[\n;]+", raw)
    else:
        parts = list(raw)
    result: list[str] = []
    for item in parts:
        item = re.sub(r"^\s*(?:\d+[.)]|[-•])\s*", "", str(item)).strip()
        if item:
            result.append(item)
    return result


def final_note(note: str) -> str:
    note = re.sub(r"\s+", " ", note).strip()
    phrase = "sebagaimana berkas terlampir"
    if phrase not in note.lower():
        note = note.rstrip(" .,:;") + " " + phrase
    return note.rstrip(".") + "."


def safe_filename(value: str) -> str:
    value = re.sub(r"[^A-Za-z0-9._-]+", "_", value)
    return value.strip("_") or "nota_konsep"


def _register_pdf_fonts() -> tuple[str, str]:
    regular_candidates = [
        "/usr/share/fonts/truetype/croscore/Arimo-Regular.ttf",
        "/usr/share/fonts/truetype/liberation2/LiberationSans-Regular.ttf",
    ]
    bold_candidates = [
        "/usr/share/fonts/truetype/croscore/Arimo-Bold.ttf",
        "/usr/share/fonts/truetype/liberation2/LiberationSans-Bold.ttf",
    ]
    regular = next((p for p in regular_candidates if Path(p).exists()), None)
    bold = next((p for p in bold_candidates if Path(p).exists()), None)
    if regular and bold:
        if "ArialCompat" not in pdfmetrics.getRegisteredFontNames():
            pdfmetrics.registerFont(TTFont("ArialCompat", regular))
        if "ArialCompat-Bold" not in pdfmetrics.getRegisteredFontNames():
            pdfmetrics.registerFont(TTFont("ArialCompat-Bold", bold))
        return "ArialCompat", "ArialCompat-Bold"
    return "Helvetica", "Helvetica-Bold"


def _draw_centered(c: canvas.Canvas, text: str, top_baseline: float, font: str, size: float, underline=False):
    page_w, page_h = A4
    y = page_h - top_baseline
    c.setFont(font, size)
    c.drawCentredString(page_w / 2, y, text)
    if underline:
        width = pdfmetrics.stringWidth(text, font, size)
        c.setLineWidth(0.7)
        c.line((page_w - width) / 2, y - 2, (page_w + width) / 2, y - 2)


def _draw_para_top(c, text, x, top, width, style):
    para = Paragraph(xml_escape(text).replace("\n", "<br/>"), style)
    _, height = para.wrap(width, 1000)
    para.drawOn(c, x, A4[1] - top - height)
    return height


def _generate_pdf(
    path: Path,
    letter_number: str,
    letter_date: str,
    recipients: list[str],
    topic: str,
    note: str,
    attachment: str,
) -> None:
    regular, bold = _register_pdf_fonts()
    c = canvas.Canvas(str(path), pagesize=A4)
    page_w, page_h = A4

    # Header, Arial-compatible 11 pt, mengikuti posisi contoh.
    _draw_centered(c, HEADER_LINES[0], 82.2, bold, 11)
    _draw_centered(c, HEADER_LINES[1], 94.7, bold, 11)
    _draw_centered(c, HEADER_LINES[2], 107.4, bold, 11)
    c.setLineWidth(0.75)
    c.line(160, page_h - 113.6, 459, page_h - 113.6)

    c.setFont(regular, 11)
    c.drawString(302.4, page_h - 139.2, f"Surabaya, {format_date_long_id(letter_date)}")
    c.drawString(302.4, page_h - 158.2, "Kepada")
    c.drawString(302.4, page_h - 177.2, "Yth. Ibu Kepala Dinas Kebudayaan")
    c.drawString(323.8, page_h - 196.2, "dan Pariwisata Provinsi Jawa Timur")
    c.drawString(323.8, page_h - 215.2, "di")
    c.drawString(339.1, page_h - 234.2, "SURABAYA")
    c.drawString(77.3, page_h - 179.2, f"Nomor: {letter_number}")

    _draw_centered(c, "NOTA PENGAJUAN KONSEP NASKAH DINAS", 291.2, bold, 11, underline=True)

    normal_style = ParagraphStyle(
        "Normal11",
        fontName=regular,
        fontSize=11,
        leading=18.96,
        alignment=TA_LEFT,
        spaceBefore=0,
        spaceAfter=0,
    )
    justify_style = ParagraphStyle(
        "Justify11",
        parent=normal_style,
        alignment=TA_JUSTIFY,
    )

    _draw_para_top(c, "Disampaikan dengan hormat :", 72, 317.0, 250, normal_style)

    label_x = 72.0
    colon_x = 144.0
    content_x = 156.0
    content_w = 390.0
    current_top = 336.0

    # Kepada.
    _draw_para_top(c, "Kepada", label_x, current_top, 65, normal_style)
    _draw_para_top(c, ":", colon_x, current_top, 10, normal_style)
    shown = recipients if len(recipients) <= 2 else ["Terlampir"]
    recipient_lines = []
    if len(recipients) <= 2:
        recipient_lines = [f"{idx}. {name}" for idx, name in enumerate(shown, start=1)]
    else:
        recipient_lines = shown
    recipient_text = "\n".join(recipient_lines)
    recipient_height = _draw_para_top(c, recipient_text, content_x, current_top, content_w, normal_style)
    current_top += max(18.96, recipient_height)

    # Dari.
    _draw_para_top(c, "Dari", label_x, current_top, 65, normal_style)
    _draw_para_top(c, ":", colon_x, current_top, 10, normal_style)
    from_h = _draw_para_top(c, FROM_TEXT, content_x, current_top, content_w, normal_style)
    current_top += max(18.96, from_h)

    # Tentang.
    _draw_para_top(c, "Tentang", label_x, current_top, 65, normal_style)
    _draw_para_top(c, ":", colon_x, current_top, 10, normal_style)
    topic_h = _draw_para_top(c, topic, content_x, current_top, content_w, normal_style)
    current_top += max(18.96, topic_h)

    # Catatan.
    _draw_para_top(c, "Catatan", label_x, current_top, 65, normal_style)
    _draw_para_top(c, ":", colon_x, current_top, 10, normal_style)
    note_h = _draw_para_top(c, final_note(note), content_x, current_top, content_w, justify_style)
    current_top += max(18.96, note_h)

    # Lampiran.
    _draw_para_top(c, "Lampiran", label_x, current_top, 65, normal_style)
    _draw_para_top(c, ":", colon_x, current_top, 10, normal_style)
    att_h = _draw_para_top(c, attachment, content_x, current_top, content_w, normal_style)
    current_top += max(18.96, att_h)

    closing_h = _draw_para_top(c, CLOSING_TEXT, label_x, current_top, 475, normal_style)
    current_top += closing_h

    if current_top > 548:
        raise ValueError(
            "Isi nota terlalu panjang untuk format satu halaman. Ringkas bagian Kepada, Tentang, atau Catatan."
        )

    # Blok bawah tetap seperti contoh.
    c.setFont(regular, 11)
    c.drawString(77.3, page_h - 575.8, "DISPOSISI PIMPINAN")
    c.drawCentredString(409.5, page_h - 575.8, SIGNATURE_TITLE[0])
    c.drawCentredString(409.5, page_h - 594.8, SIGNATURE_TITLE[1])

    c.setFont(bold, 11)
    name_y = page_h - 655.5
    c.drawCentredString(409.5, name_y, SIGNATURE_NAME)
    name_width = pdfmetrics.stringWidth(SIGNATURE_NAME, bold, 11)
    c.setLineWidth(0.65)
    c.line(409.5 - name_width / 2, name_y - 2, 409.5 + name_width / 2, name_y - 2)
    c.setFont(regular, 11)
    c.drawCentredString(409.5, page_h - 668.2, SIGNATURE_RANK)
    c.drawCentredString(409.5, page_h - 681.0, SIGNATURE_NIP)

    c.showPage()
    c.save()




def _set_table_fixed_widths(table, widths):
    table.autofit = False
    tblPr = table._tbl.tblPr
    layout = tblPr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tblPr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid_cols = table._tbl.tblGrid.gridCol_lst
    for index, width in enumerate(widths):
        if index < len(grid_cols):
            grid_cols[index].set(qn("w:w"), str(int(width.twips)))
        for row in table.rows:
            cell = row.cells[index]
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.first_child_found_in("w:tcW")
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.append(tcW)
            tcW.set(qn("w:w"), str(int(width.twips)))
            tcW.set(qn("w:type"), "dxa")


def _set_cell_margins(cell, top=0, start=0, bottom=0, end=0):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def _remove_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = tblPr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblPr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "nil")


def _set_run_arial(run, bold: bool | None = None, underline: bool | None = None):
    run.font.name = "Arial"
    run.font.size = Pt(11)
    if bold is not None:
        run.bold = bold
    if underline is not None:
        run.underline = underline
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.rFonts
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rFonts.set(qn(f"w:{attr}"), "Arial")


def _format_para(p, alignment=WD_ALIGN_PARAGRAPH.LEFT, line=Pt(19), before=0, after=0):
    p.alignment = alignment
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = line


def _add_text(p, text, bold=False, underline=False):
    run = p.add_run(text)
    _set_run_arial(run, bold=bold, underline=underline)
    return run


def _add_bottom_border(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _generate_docx(
    path: Path,
    letter_number: str,
    letter_date: str,
    recipients: list[str],
    topic: str,
    note: str,
    attachment: str,
) -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(1.7)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")

    header = doc.add_paragraph()
    _format_para(header, WD_ALIGN_PARAGRAPH.CENTER, Pt(12.5), after=0)
    for idx, line in enumerate(HEADER_LINES):
        _add_text(header, line, bold=True)
        if idx < len(HEADER_LINES) - 1:
            header.add_run().add_break()
    _add_bottom_border(header)

    # Area nomor dan alamat kepala dinas.
    top_table = doc.add_table(rows=1, cols=2)
    top_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    top_table.autofit = False
    _remove_table_borders(top_table)
    _set_table_fixed_widths(top_table, [Cm(8.1), Cm(8.6)])
    for cell in top_table.rows[0].cells:
        _set_cell_margins(cell, top=0, start=0, bottom=0, end=0)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    left = top_table.cell(0, 0)
    p = left.paragraphs[0]
    _format_para(p, line=Pt(19), before=44)
    _add_text(p, f"Nomor: {letter_number}")

    right = top_table.cell(0, 1)
    right.text = ""
    lines = [
        f"Surabaya, {format_date_long_id(letter_date)}",
        "Kepada",
        "Yth. Ibu Kepala Dinas Kebudayaan",
        "    dan Pariwisata Provinsi Jawa Timur",
        "    di",
        "        SURABAYA",
    ]
    for idx, line in enumerate(lines):
        p = right.add_paragraph() if idx else right.paragraphs[0]
        _format_para(p, line=Pt(19))
        _add_text(p, line)

    title = doc.add_paragraph()
    _format_para(title, WD_ALIGN_PARAGRAPH.CENTER, Pt(19), before=24, after=22)
    _add_text(title, "NOTA PENGAJUAN KONSEP NASKAH DINAS", bold=True, underline=True)

    intro = doc.add_paragraph()
    _format_para(intro, line=Pt(19))
    _add_text(intro, "Disampaikan dengan hormat :")

    body = doc.add_table(rows=0, cols=3)
    body.alignment = WD_TABLE_ALIGNMENT.CENTER
    body.autofit = False
    _remove_table_borders(body)
    widths = [Cm(2.45), Cm(0.45), Cm(13.8)]
    _set_table_fixed_widths(body, widths)

    def add_row(label: str, content_lines: list[str], justify=False):
        row = body.add_row()
        _set_table_fixed_widths(body, widths)
        for idx, cell in enumerate(row.cells):
            _set_cell_margins(cell, top=0, start=0, bottom=0, end=0)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        p0 = row.cells[0].paragraphs[0]
        _format_para(p0, line=Pt(19))
        _add_text(p0, label)
        p1 = row.cells[1].paragraphs[0]
        _format_para(p1, line=Pt(19))
        _add_text(p1, ":")
        row.cells[2].text = ""
        for idx, line in enumerate(content_lines):
            p2 = row.cells[2].add_paragraph() if idx else row.cells[2].paragraphs[0]
            _format_para(
                p2,
                WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT,
                Pt(19),
            )
            _add_text(p2, line)

    if len(recipients) <= 2:
        shown_recipients = [f"{i}. {value}" for i, value in enumerate(recipients, 1)]
    else:
        shown_recipients = ["Terlampir"]

    add_row("Kepada", shown_recipients)
    add_row("Dari", [FROM_TEXT])
    add_row("Tentang", [topic])
    add_row("Catatan", [final_note(note)], justify=True)
    add_row("Lampiran", [attachment])

    closing = doc.add_paragraph()
    _format_para(closing, line=Pt(19))
    _add_text(closing, CLOSING_TEXT)

    bottom = doc.add_table(rows=1, cols=2)
    bottom.alignment = WD_TABLE_ALIGNMENT.CENTER
    bottom.autofit = False
    _remove_table_borders(bottom)
    _set_table_fixed_widths(bottom, [Cm(7.6), Cm(9.1)])
    for cell in bottom.rows[0].cells:
        _set_cell_margins(cell, top=0, start=0, bottom=0, end=0)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    p = bottom.cell(0, 0).paragraphs[0]
    _format_para(p, line=Pt(19), before=27)
    _add_text(p, "DISPOSISI PIMPINAN")

    sig = bottom.cell(0, 1)
    sig.text = ""
    p = sig.paragraphs[0]
    _format_para(p, WD_ALIGN_PARAGRAPH.CENTER, Pt(19), before=27)
    _add_text(p, SIGNATURE_TITLE[0])
    p = sig.add_paragraph()
    _format_para(p, WD_ALIGN_PARAGRAPH.CENTER, Pt(19))
    _add_text(p, SIGNATURE_TITLE[1])
    p = sig.add_paragraph()
    _format_para(p, WD_ALIGN_PARAGRAPH.CENTER, Pt(19), before=38)
    _add_text(p, SIGNATURE_NAME, bold=True, underline=True)
    p = sig.add_paragraph()
    _format_para(p, WD_ALIGN_PARAGRAPH.CENTER, Pt(12.5))
    _add_text(p, SIGNATURE_RANK)
    p = sig.add_paragraph()
    _format_para(p, WD_ALIGN_PARAGRAPH.CENTER, Pt(12.5))
    _add_text(p, SIGNATURE_NIP)

    # Pastikan seluruh run memakai Arial 11.
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            _set_run_arial(run)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        _set_run_arial(run)

    doc.save(path)


def generate_nota_files(
    letter_number: str,
    letter_date: str,
    recipients: str | Iterable[str],
    topic: str,
    note: str,
    attachment: str,
    output_dir: str | Path | None = None,
) -> tuple[str, str]:
    recipients_list = clean_recipients(recipients)
    if not recipients_list:
        raise ValueError("Daftar Kepada tidak boleh kosong.")
    if len(topic.strip()) < 3:
        raise ValueError("Tentang terlalu pendek.")
    if len(topic) > 220:
        raise ValueError("Tentang terlalu panjang. Maksimal 220 karakter.")
    if len(note.strip()) < 5:
        raise ValueError("Catatan terlalu pendek.")
    if len(note) > 650:
        raise ValueError("Catatan terlalu panjang. Maksimal 650 karakter agar format tetap satu halaman.")
    if attachment not in {"-", "1 (satu) berkas", "2 (dua) berkas"}:
        raise ValueError("Pilihan lampiran tidak valid.")

    if output_dir is None:
        output_dir = tempfile.mkdtemp(prefix="nota_konsep_")
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    stem = safe_filename(f"Nota_Konsep_{letter_number}")
    docx_path = output_dir / f"{stem}.docx"
    pdf_path = output_dir / f"{stem}.pdf"

    _generate_docx(docx_path, letter_number, letter_date, recipients_list, topic.strip(), note.strip(), attachment)
    _generate_pdf(pdf_path, letter_number, letter_date, recipients_list, topic.strip(), note.strip(), attachment)
    return str(docx_path), str(pdf_path)
